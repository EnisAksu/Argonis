import sys
import os
import json
import tkinter as tk
from tkinter import ttk, messagebox
import subprocess
from datetime import datetime

# Ensure dependencies are installed
def install_and_import(package):
    try:
        __import__(package)
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])
    finally:
        globals()[package] = __import__(package)

# Install required packages
install_and_import("customtkinter")
install_and_import("darkdetect")

import customtkinter as ctk

# Set appearance mode based on system settings
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

class NIS2ComplianceTool:
    def __init__(self):
        self.window = ctk.CTk()
        self.window.title("NIS2 Compliance Tool")
        self.window.geometry("1200x800")
        
        # Configure grid
        self.window.grid_columnconfigure(0, weight=1)
        self.window.grid_rowconfigure(0, weight=1)
        
        # Create main container
        self.main_container = ctk.CTkFrame(self.window)
        self.main_container.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        
        # Create tab view
        self.tab_view = ctk.CTkTabview(self.main_container)
        self.tab_view.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        
        # Configure main container grid
        self.main_container.grid_columnconfigure(0, weight=1)
        self.main_container.grid_rowconfigure(0, weight=1)
        
        # Create tabs
        self.checker_tab = self.tab_view.add("Compliance Checker")
        self.guidance_tab = self.tab_view.add("Compliance Guidance")
        
        # Initialize tabs
        self.setup_checker_tab()
        self.setup_guidance_tab()
        
        # Load saved data if exists
        self.load_saved_data()

    def setup_checker_tab(self):
        # Organization Size Frame
        self.size_frame = ctk.CTkFrame(self.checker_tab)
        self.size_frame.pack(fill="x", padx=10, pady=5)
        
        ctk.CTkLabel(self.size_frame, text="Organization Size").pack()
        
        # Staff headcount
        self.staff_var = tk.StringVar()
        self.staff_dropdown = ctk.CTkOptionMenu(
            self.size_frame,
            variable=self.staff_var,
            values=["< 10", "10-49", "50-249", "> 250"]
        )
        self.staff_dropdown.pack(pady=5)
        
        # Turnover
        self.turnover_var = tk.StringVar()
        self.turnover_dropdown = ctk.CTkOptionMenu(
            self.size_frame,
            variable=self.turnover_var,
            values=["< €2M", "€2M-€10M", "€10M-€50M", "> €50M"]
        )
        self.turnover_dropdown.pack(pady=5)

        # Sector Selection Frame
        self.sector_frame = ctk.CTkFrame(self.checker_tab)
        self.sector_frame.pack(fill="x", padx=10, pady=5)
        
        ctk.CTkLabel(self.sector_frame, text="Sector").pack()

        # Create a frame for the three columns
        columns_frame = ctk.CTkFrame(self.sector_frame)
        columns_frame.pack(fill="x", padx=5, pady=5)

        # Configure the columns to be equal width
        for i in range(3):
            columns_frame.grid_columnconfigure(i, weight=1)

        self.sectors = [
            "Energy",
            "Transport",
            "Banking",
            "Financial Market Infrastructure",
            "Healthcare",
            "Drinking Water",
            "Digital Infrastructure",
            "Public Administration",
            "Space",
            "Postal Services",
            "Waste Management",
            "Digital Services",
            "Manufacturing",
            "Food Production",
            "None of the above"
        ]

        # Calculate items per column
        items_per_column = -(-len(self.sectors) // 3)  # Ceiling division

        self.sector_var = tk.StringVar()
        for i, sector in enumerate(self.sectors):
            column = i // items_per_column
            row = i % items_per_column
            radio_btn = ctk.CTkRadioButton(
                columns_frame,
                text=sector,
                variable=self.sector_var,
                value=sector
            )
            radio_btn.grid(row=row, column=column, padx=5, pady=2, sticky="w")

        # Check Compliance Button
        self.check_button = ctk.CTkButton(
            self.checker_tab,
            text="Check Compliance Status",
            command=self.check_compliance
        )
        self.check_button.pack(pady=20)

        # Results Frame to hold both text and save button
        self.results_frame = ctk.CTkFrame(self.checker_tab)
        self.results_frame.pack(pady=10, padx=10, fill="both", expand=True)

        # Results Display
        self.results_text = ctk.CTkTextbox(
            self.results_frame,
            height=200,
            width=500
        )
        self.results_text.pack(pady=10, padx=10, fill="both", expand=True)

        # Save Results Button (hidden initially)
        self.save_results_button = ctk.CTkButton(
            self.results_frame,
            text="Save the output",
            command=self.save_compliance_results,
            width=120
        )
        self.save_results_button.pack(side="left", padx=10, pady=(0, 10))
        self.save_results_button.pack_forget()

    def setup_guidance_tab(self):
        # Company Information Frame
        self.info_frame = ctk.CTkFrame(self.guidance_tab)
        self.info_frame.pack(fill="x", padx=10, pady=5)
        
        # Company Name
        ctk.CTkLabel(self.info_frame, text="Company Name").pack()
        self.company_name = ctk.CTkEntry(self.info_frame)
        self.company_name.pack(pady=5)
        
        # Industry Sector (using the same sectors as checker)
        ctk.CTkLabel(self.info_frame, text="Industry Sector").pack()
        self.guidance_sector_var = tk.StringVar()
        self.guidance_sector_dropdown = ctk.CTkOptionMenu(
            self.info_frame,
            variable=self.guidance_sector_var,
            values=self.sectors
        )
        self.guidance_sector_dropdown.pack(pady=5)
        
        # Number of Employees
        ctk.CTkLabel(self.info_frame, text="Number of Employees").pack()
        self.employee_count = ctk.CTkEntry(self.info_frame)
        self.employee_count.pack(pady=5)
        
        # Annual Revenue
        ctk.CTkLabel(self.info_frame, text="Annual Revenue (€)").pack()
        self.revenue = ctk.CTkEntry(self.info_frame)
        self.revenue.pack(pady=5)
        
        # Create Checklist Button
        self.create_checklist_button = ctk.CTkButton(
            self.guidance_tab,
            text="Create My Checklist",
            command=self.create_checklist
        )
        self.create_checklist_button.pack(pady=20)
        
        # Results frame
        self.checklist_frame = ctk.CTkFrame(self.guidance_tab)
        self.checklist_frame.pack(fill="both", expand=True, padx=10, pady=5)
        
        # Checklist Display
        self.checklist_text = ctk.CTkTextbox(
            self.checklist_frame,
            height=400,
            width=500
        )
        self.checklist_text.pack(fill="both", expand=True, padx=10, pady=(10, 5))
        
        # Save button frame at the bottom of checklist frame
        self.save_button_frame = ctk.CTkFrame(self.checklist_frame)
        self.save_button_frame.pack(fill="x", side="bottom", padx=10, pady=5)
        
        # Save Button
        self.save_checklist_button = ctk.CTkButton(
            self.save_button_frame,
            text="Save the output",
            command=self.save_checklist_docx,
            width=120
        )
        self.save_checklist_button.pack(side="left", padx=10, pady=5)
        self.save_checklist_button.pack_forget()

    def check_compliance(self):
        # Clear previous results
        self.results_text.delete("1.0", tk.END)
        
        # Get values
        staff = self.staff_var.get()
        turnover = self.turnover_var.get()
        sector = self.sector_var.get()
        
        # Determine compliance status
        is_essential = False
        is_important = False
        
        # Essential entity criteria
        if sector in ["Energy", "Transport", "Banking", "Financial Market Infrastructure",
                     "Healthcare", "Drinking Water", "Digital Infrastructure"]:
            if staff in ["50-249", "> 250"] or turnover in ["€10M-€50M", "> €50M"]:
                is_essential = True
        
        # Important entity criteria
        elif sector not in ["None of the above"]:
            if staff in ["10-49", "50-249"] or turnover in ["€2M-€10M", "€10M-€50M"]:
                is_important = True
        
        # Display results
        if is_essential:
            self.results_text.insert(tk.END, "Your organization is classified as an ESSENTIAL entity under NIS2.\n\n")
            self.results_text.insert(tk.END, "Requirements include:\n")
            self.results_text.insert(tk.END, "- Implementing comprehensive cybersecurity risk management measures\n")
            self.results_text.insert(tk.END, "- Mandatory incident reporting within 24 hours\n")
            self.results_text.insert(tk.END, "- Regular security audits and assessments\n")
            self.results_text.insert(tk.END, "- Designation of CISO or equivalent role\n")
        elif is_important:
            self.results_text.insert(tk.END, "Your organization is classified as an IMPORTANT entity under NIS2.\n\n")
            self.results_text.insert(tk.END, "Requirements include:\n")
            self.results_text.insert(tk.END, "- Basic cybersecurity risk management measures\n")
            self.results_text.insert(tk.END, "- Incident reporting for significant events\n")
            self.results_text.insert(tk.END, "- Periodic security reviews\n")
        else:
            self.results_text.insert(tk.END, "Your organization appears to be OUT OF SCOPE of NIS2.\n\n")
            self.results_text.insert(tk.END, "However, implementing cybersecurity best practices is still recommended.\n")

        # After displaying results, show the save button
        self.save_results_button.pack(side="left", padx=10, pady=(0, 10))

    def create_checklist(self):
        # Clear previous checklist
        self.checklist_text.delete("1.0", tk.END)
        
        # Get company information
        company_name = self.company_name.get()
        sector = self.guidance_sector_var.get()
        employees = self.employee_count.get()
        revenue = self.revenue.get()
        
        # Generate checklist header
        self.checklist_text.insert(tk.END, f"NIS2 Compliance Checklist for {company_name}\n")
        self.checklist_text.insert(tk.END, f"Generated on: {datetime.now().strftime('%Y-%m-%d')}\n\n")
        
        # Basic requirements for all entities
        self.checklist_text.insert(tk.END, "1. Basic Cybersecurity Measures:\n")
        self.checklist_text.insert(tk.END, "□ Implement risk assessment procedures\n")
        self.checklist_text.insert(tk.END, "□ Establish incident handling processes\n")
        self.checklist_text.insert(tk.END, "□ Deploy business continuity measures\n")
        self.checklist_text.insert(tk.END, "□ Implement network security measures\n\n")
        
        # Sector-specific requirements
        self.checklist_text.insert(tk.END, "2. Sector-Specific Requirements:\n")
        if sector == "Energy":
            self.checklist_text.insert(tk.END, "□ Implement SCADA system security\n")
            self.checklist_text.insert(tk.END, "□ Establish energy-specific incident response\n")
        elif sector == "Healthcare":
            self.checklist_text.insert(tk.END, "□ Implement medical device security\n")
            self.checklist_text.insert(tk.END, "□ Establish patient data protection measures\n")
        
        # Size-based requirements
        self.checklist_text.insert(tk.END, "\n3. Organization-Specific Requirements:\n")
        try:
            emp_count = int(employees)
            if emp_count > 250:
                self.checklist_text.insert(tk.END, "□ Appoint dedicated CISO\n")
                self.checklist_text.insert(tk.END, "□ Establish security operations center\n")
            elif emp_count > 50:
                self.checklist_text.insert(tk.END, "□ Designate security responsible person\n")
                self.checklist_text.insert(tk.END, "□ Implement basic security monitoring\n")
        except ValueError:
            pass
        
        # Compliance reporting requirements
        self.checklist_text.insert(tk.END, "\n4. Reporting Requirements:\n")
        self.checklist_text.insert(tk.END, "□ Establish incident reporting procedures\n")
        self.checklist_text.insert(tk.END, "□ Create compliance documentation process\n")
        self.checklist_text.insert(tk.END, "□ Set up regular reporting schedule\n")

        # Show the save button
        self.save_checklist_button.pack(side="left", padx=10, pady=5)

    def save_compliance_results(self):
        try:
            with open("Argonis Compliance Checker.TXT", "w") as f:
                f.write(self.results_text.get("1.0", tk.END))
            messagebox.showinfo("Success", "Results saved to 'Argonis Compliance Checker.TXT'")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save results: {str(e)}")

    def save_checklist_docx(self):
        try:
            # First ensure python-docx is installed
            subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
            from docx import Document
            from docx.shared import Pt
        
            company_name = self.company_name.get()
            filename = f"NIS2_Checklist_{company_name}_{datetime.now().strftime('%Y%m%d')}.docx"
        
            # Create a new Word document
            doc = Document()
            
            # Get checklist content
            checklist_content = self.checklist_text.get("1.0", tk.END)
            
            # Add title with company name
            doc.add_heading(f'NIS2 Compliance Checklist for {company_name}', 0)
            
            # Add generation date
            doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d")}')
            doc.add_paragraph()  # Add blank line
            
            # Process each line
            current_section = None
            for line in checklist_content.split('\n'):
                if line.strip():
                    if line.endswith(':'):  # Section header
                        doc.add_heading(line, level=1)
                    elif line.startswith('□'):  # Checklist item
                        # Add checkbox and text as a paragraph
                        p = doc.add_paragraph()
                        p.add_run('☐ ').font.name = 'Segoe UI Symbol'  # Using a different checkbox symbol
                        p.add_run(line[2:].strip())
            
            # Save the document
            doc.save(filename)
            messagebox.showinfo("Success", f"Checklist saved as {filename}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save checklist: {str(e)}")

    def load_saved_data(self):
        try:
            with open("nis2_compliance_data.json", "r") as f:
                data = json.load(f)
                # Load saved values into fields
                self.staff_var.set(data.get("staff", ""))
                self.turnover_var.set(data.get("turnover", ""))
                self.sector_var.set(data.get("sector", ""))
                self.company_name.insert(0, data.get("company_name", ""))
                self.guidance_sector_var.set(data.get("guidance_sector", ""))
                self.employee_count.insert(0, data.get("employees", ""))
                self.revenue.insert(0, data.get("revenue", ""))
        except FileNotFoundError:
            pass

    def save_data(self):
        data = {
            "staff": self.staff_var.get(),
            "turnover": self.turnover_var.get(),
            "sector": self.sector_var.get(),
            "company_name": self.company_name.get(),
            "guidance_sector": self.guidance_sector_var.get(),
            "employees": self.employee_count.get(),
            "revenue": self.revenue.get()
        }
        with open("nis2_compliance_data.json", "w") as f:
            json.dump(data, f)

    def run(self):
        self.window.mainloop()

if __name__ == "__main__":
    app = NIS2ComplianceTool()
    app.run()